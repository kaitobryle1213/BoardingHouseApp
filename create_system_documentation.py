#!/usr/bin/env python3
"""
Create comprehensive system documentation for Boarding House Management System
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def create_system_documentation():
    """Create comprehensive system documentation"""
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Boarding House Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_heading('Complete System Documentation', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', 1)
    
    toc_items = [
        '1. System Overview',
        '2. Database Models',
        '3. User Management',
        '4. Room Management',
        '5. Customer Management',
        '6. Payment System',
        '7. Monthly Cycle System',
        '8. API Endpoints',
        '9. Forms and Validation',
        '10. Templates and UI',
        '11. Security Features',
        '12. Business Rules',
        '13. Installation and Setup',
        '14. Troubleshooting'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 1: System Overview
    doc.add_heading('1. System Overview', 1)
    doc.add_paragraph(
        'The Boarding House Management System is a Django-based web application designed to manage '
        'boarding house operations including room management, customer tracking, payment processing, '
        'and financial reporting.'
    )
    
    doc.add_heading('Key Features', 2)
    features = [
        'User authentication and role-based access control',
        'Room management with capacity tracking',
        'Customer management with detailed profiles',
        'Monthly payment cycle system',
        'Payment processing with receipt generation',
        'Room transfer management',
        'Financial reporting and analytics',
        'Responsive web interface'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 2: Database Models
    doc.add_heading('2. Database Models', 1)
    
    doc.add_heading('BoardingHouseUser Model', 2)
    doc.add_paragraph('Extends Django AbstractUser with additional fields:')
    user_fields = [
        'role: Admin/User role selection',
        'status: Active/Inactive status',
        'date_created: Account creation timestamp'
    ]
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Room Model', 2)
    doc.add_paragraph('Manages room information and occupancy:')
    room_fields = [
        'room_number: Unique room identifier',
        'status: Available/Under Maintenance/Occupied',
        'room_type: Single/Bed Spacer',
        'capacity: Maximum occupancy',
        'price: Monthly rental price',
        'date_created: Room creation date',
        'date_left: Room vacation date'
    ]
    for field in room_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Customer Model', 2)
    doc.add_paragraph('Stores customer information and payment details:')
    customer_fields = [
        'customer_id: Auto-generated unique ID',
        'name: Customer full name',
        'address: Complete address',
        'contact_number: Phone number',
        'parents_name: Emergency contact',
        'parents_contact_number: Emergency phone',
        'status: Active/Inactive status',
        'room: Foreign key to Room',
        'due_date: Monthly payment due date',
        'date_entry: Move-in date',
        'date_left: Move-out date'
    ]
    for field in customer_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Payment Model', 2)
    doc.add_paragraph('Tracks all payment transactions:')
    payment_fields = [
        'customer: Foreign key to Customer',
        'due_date: Payment cycle date',
        'previous_date: Previous due date (for reference)',
        'amount: Expected payment amount',
        'amount_received: Actual amount received',
        'change_amount: Change given',
        'is_paid: Payment status',
        'date_paid: Payment timestamp',
        'remarks: Additional notes'
    ]
    for field in payment_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('RoomTransferHistory Model', 2)
    doc.add_paragraph('Tracks room transfer history:')
    transfer_fields = [
        'customer: Foreign key to Customer',
        'room_from: Previous room',
        'room_to: New room',
        'room_from_price: Price at time of transfer',
        'room_to_price: New room price'
    ]
    for field in transfer_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 3: User Management
    doc.add_heading('3. User Management', 1)
    doc.add_paragraph('The system supports two user roles:')
    
    doc.add_heading('Admin Users', 2)
    admin_capabilities = [
        'Full access to all system features',
        'User management (create, edit, delete users)',
        'Room management',
        'Customer management',
        'Payment processing',
        'Report generation',
        'System configuration'
    ]
    for capability in admin_capabilities:
        doc.add_paragraph(capability, style='List Bullet')
    
    doc.add_heading('Regular Users', 2)
    user_capabilities = [
        'View dashboard information',
        'Access customer information (read-only)',
        'View room occupancy',
        'Limited system access'
    ]
    for capability in user_capabilities:
        doc.add_paragraph(capability, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 4: Room Management
    doc.add_heading('4. Room Management', 1)
    
    doc.add_heading('Room Types', 2)
    room_types = [
        'Single Rooms: Individual rooms with capacity 1',
        'Bed Spacer Rooms: Shared rooms with configurable capacity'
    ]
    for room_type in room_types:
        doc.add_paragraph(room_type, style='List Bullet')
    
    doc.add_heading('Room Status', 2)
    status_types = [
        'Available: Ready for occupancy',
        'Occupied: Currently occupied by customers',
        'Under Maintenance: Not available for occupancy'
    ]
    for status in status_types:
        doc.add_paragraph(status, style='List Bullet')
    
    doc.add_heading('Room Operations', 2)
    operations = [
        'Room creation with validation',
        'Room editing and updates',
        'Room deletion with occupant transfer',
        'Occupancy tracking and status automation',
        'Capacity-based room management'
    ]
    for operation in operations:
        doc.add_paragraph(operation, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 5: Customer Management
    doc.add_heading('5. Customer Management', 1)
    
    doc.add_heading('Customer Lifecycle', 2)
    lifecycle = [
        'Customer registration with complete profile',
        'Room assignment and due date setting',
        'Active status management',
        'Payment tracking and history',
        'Room transfers (if needed)',
        'Move-out processing and status update'
    ]
    for stage in lifecycle:
        doc.add_paragraph(stage, style='List Bullet')
    
    doc.add_heading('Customer Information', 2)
    info_points = [
        'Personal details and contact information',
        'Emergency contact details',
        'Room assignment history',
        'Payment history and records',
        'Status tracking (Active/Inactive)'
    ]
    for point in info_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 6: Payment System
    doc.add_heading('6. Payment System', 1)
    
    doc.add_heading('Payment Processing', 2)
    payment_steps = [
        'Customer selection and balance checking',
        'Payment amount entry',
        'Change calculation',
        'Receipt generation',
        'Cycle status update',
        'Due date advancement (if cycle completed)'
    ]
    for step in payment_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('Payment Types', 2)
    payment_types = [
        'Full payment: Complete monthly amount',
        'Partial payment: Multiple payments per cycle',
        'Overpayment: Credit applied to next cycle',
        'Room transfer adjustments: Price difference handling'
    ]
    for p_type in payment_types:
        doc.add_paragraph(p_type, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 7: Monthly Cycle System
    doc.add_heading('7. Monthly Cycle System', 1)
    
    doc.add_paragraph('The system implements a sophisticated monthly billing cycle:')
    
    cycle_features = [
        'Automatic due date management',
        'Cycle-based payment tracking',
        'Room price-based amount calculation',
        'Automatic due date advancement',
        'Status-based color coding',
        'Overpayment handling and credit application',
        'Room transfer mid-cycle adjustments'
    ]
    for feature in cycle_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Status Colors', 2)
    status_colors = [
        'Green: Paid today',
        'Red: Due today',
        'Yellow: Due within 3 days',
        'Black: Overdue',
        'White: Upcoming or no schedule'
    ]
    for color in status_colors:
        doc.add_paragraph(color, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 8: API Endpoints
    doc.add_heading('8. API Endpoints', 1)
    
    endpoints = [
        ('/dashboard/', 'GET', 'Main dashboard view'),
        ('/api/dashboard_data/', 'GET', 'Dashboard data API'),
        ('/api/customers_data/', 'GET', 'Customers data API'),
        ('/users/', 'GET', 'User management'),
        ('/users/create/', 'POST', 'Create new user'),
        ('/users/<pk>/edit/', 'POST', 'Edit user'),
        ('/customers/', 'GET', 'Customer management'),
        ('/customers/create/', 'POST', 'Create customer'),
        ('/customers/<pk>/edit/', 'POST', 'Edit customer'),
        ('/customers/<pk>/delete/', 'POST', 'Delete customer'),
        ('/customers/<pk>/', 'GET', 'Customer details'),
        ('/payment/', 'GET/POST', 'Payment processing'),
        ('/api/search_customers/', 'GET', 'Customer search'),
        ('/api/get_balance/<customer_id>/', 'GET', 'Get customer balance'),
        ('/api/process_payment/', 'POST', 'Process payment'),
        ('/rooms/', 'GET', 'Room management'),
        ('/rooms/create/', 'POST', 'Create room'),
        ('/rooms/<pk>/edit/', 'POST', 'Edit room'),
        ('/rooms/<pk>/delete/', 'POST', 'Delete room'),
        ('/rooms/<pk>/occupants/', 'GET', 'Room occupants'),
        ('/rooms/transfer/<customer_id>/', 'POST', 'Transfer customer'),
        ('/api/search_rooms/', 'GET', 'Room search'),
        ('/report/', 'GET', 'Financial reports'),
        ('/report/transfers/', 'GET', 'Transfer reports'),
        ('/logout/', 'GET', 'User logout')
    ]
    
    for endpoint, method, description in endpoints:
        doc.add_paragraph(f'{method} {endpoint} - {description}')
    
    doc.add_page_break()
    
    # Section 9: Forms and Validation
    doc.add_heading('9. Forms and Validation', 1)
    
    doc.add_heading('RoomForm', 2)
    room_form = [
        'Validates room number uniqueness',
        'Enforces capacity rules based on room type',
        'Single rooms: Capacity must be 1',
        'Bed spacer rooms: Minimum capacity 1',
        'Price validation and formatting'
    ]
    for validation in room_form:
        doc.add_paragraph(validation, style='List Bullet')
    
    doc.add_heading('CustomerForm', 2)
    customer_form = [
        'Personal information validation',
        'Contact number formatting',
        'Room assignment validation',
        'Due date setting and validation'
    ]
    for validation in customer_form:
        doc.add_paragraph(validation, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 10: Templates and UI
    doc.add_heading('10. Templates and UI', 1)
    
    templates = [
        'base.html: Base template with navigation',
        'dashboard.html: Main dashboard with customer status',
        'login.html: User authentication',
        'customer.html: Customer list and management',
        'customer_form.html: Customer creation/editing',
        'customer_detail.html: Customer details view',
        'payment.html: Payment processing interface',
        'room.html: Room management',
        'room_form.html: Room creation/editing',
        'room_occupants.html: Room occupancy view',
        'room_delete.html: Room deletion confirmation',
        'report.html: Financial reports',
        'transfer_report.html: Transfer history reports',
        'user.html: User management',
        'user_form.html: User creation/editing'
    ]
    
    for template in templates:
        doc.add_paragraph(template, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 11: Security Features
    doc.add_heading('11. Security Features', 1)
    
    security = [
        'Django authentication system',
        'Role-based access control',
        'Admin-only endpoints protection',
        'CSRF protection',
        'Form validation and sanitization',
        'SQL injection prevention',
        'XSS protection'
    ]
    
    for feature in security:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 12: Business Rules
    doc.add_heading('12. Business Rules', 1)
    
    rules = [
        'Monthly cycles correspond to calendar months',
        'Payments tracked per due date (per cycle)',
        'Status determined by paid amount vs room price',
        'Due dates advance only when cycle fully paid',
        'Room price changes handled during transfers',
        'Overpayments credited to next cycle',
        'Partial payments supported within cycles'
    ]
    
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 13: Installation and Setup
    doc.add_heading('13. Installation and Setup', 1)
    
    setup_steps = [
        'Python 3.13+ required',
        'Django 6.0 framework',
        'Virtual environment setup',
        'Database migration',
        'Superuser creation',
        'Static files collection',
        'Development server startup'
    ]
    
    for step in setup_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 14: Troubleshooting
    doc.add_heading('14. Troubleshooting', 1)
    
    issues = [
        'Due date not advancing: Check if cycle fully paid',
        'Incorrect status: Verify payment amounts match room price',
        'Room transfer problems: Ensure payment adjustment logic executed',
        'Database issues: Run migrations and check connections',
        'Permission errors: Verify user roles and access rights'
    ]
    
    for issue in issues:
        doc.add_paragraph(issue, style='List Bullet')
    
    # Save the document
    doc.save('Boarding_House_System_Documentation.docx')
    print("System documentation created successfully!")

if __name__ == "__main__":
    create_system_documentation()